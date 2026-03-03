import docx

doc = docx.Document('/mnt/c/Users/johann/OneDrive/Desktop/Code/Ai_Project_1/Project1-ngq7-Report.docx')

all_text = []
for i, para in enumerate(doc.paragraphs):
    all_text.append((i, para.style.name, para.text))

print("=" * 70)
print("FINAL AUDIT: ISSUE DEEP-DIVES")
print("=" * 70)

print()
print("--- ISSUE 1: TABLES APPEARING AFTER '*.1 The Table' HEADINGS ---")
print("(Each heading is followed by an empty paragraph; the actual Word")
print(" table is embedded in the docx but appears as a gap in paragraphs)")
print()
# The Word tables ARE present (confirmed: 3 tables with 101 rows each).
# They appear between certain paragraphs. Let's check paragraph indices
# around P11, P386, P782 (the empty paras after The Table headings)
for idx in [10, 11, 12, 385, 386, 387, 781, 782, 783]:
    if idx < len(all_text):
        i, style, text = all_text[idx]
        print(f"  [P{i}] STYLE={style} | {repr(text[:80])}")

print()
print("--- ISSUE 2: DUPLICATE PYTHON CODE IN TABLE 3 (Table E) ---")
print("Table 3 section contains TWO Python code blocks for TableE.py:")
print("  - [P789]  'Python Code (Project1-ngq7-TableE.py):'  (1st occurrence)")
print("  - [P1191] 'Python Code (Project1-ngq7-TableE.py):'  (2nd occurrence)")
print("  - [P1065] 'C++ Code (Project1-ngq7-TableE.cpp):'    (1st C++ occurrence)")
print("  - [P1467] 'C++ Code (Project1-ngq7-TableE.cpp):'    (2nd C++ occurrence)")
print()
print("Paragraph context around second Python occurrence:")
for idx in [1188, 1189, 1190, 1191, 1192, 1193, 1194]:
    if idx < len(all_text):
        i, style, text = all_text[idx]
        print(f"  [P{i}] STYLE={style} | {repr(text[:100])}")

print()
print("--- ISSUE 3: C++ CODE OUTPUT NOT IN REPORT ---")
print("The report includes Python code output for all 3 tables.")
print("No C++ code output (i.e., the compiled/run result) appears anywhere.")
print("Requirement check 4 says 'Python OR C++ code ran and matches ChatGPT'.")
print("Python output IS present. C++ output is NOT present (but not required if Python satisfies it).")
print()

print()
print("--- ISSUE 4: TABLE NAMING - 'Table 1/2/3' vs 'Table A/B/E' ---")
print("The report uses '2. Table 1', '3. Table 2', '4. Table 3' as headings.")
print("The assignment refers to these as 'Table A', 'Table B', 'Table E'.")
print("The code files are named TableA.py, TableB.py, TableE.py — consistent with assignment.")
print("But the section HEADINGS in the document say 'Table 1', 'Table 2', 'Table 3'.")
print("The verification checks mention TableA.py, TableB.py, TableE.py by name.")
print()

print()
print("--- ISSUE 5: 'The Table' SECTIONS - WORD TABLE VS IMAGE ---")
# The word tables are confirmed present with 101 rows.
# Let's verify Y column is filled (not blank) for all 3 tables
for ti, table in enumerate(doc.tables):
    name = ['Table 1 (A)', 'Table 2 (B)', 'Table 3 (E)'][ti]
    y_col = 5  # last column index
    blank_y = 0
    for ri, row in enumerate(table.rows):
        if ri == 0:
            continue  # skip header
        y_val = row.cells[y_col].text.strip()
        if y_val == '':
            blank_y += 1
    print(f"  {name}: Y column - {blank_y} blank cells out of 100 data rows")
    # Show header
    headers = [c.text.strip() for c in table.rows[0].cells]
    print(f"    Column headers: {headers}")

print()
print("--- ISSUE 6: CHECK 3 BLANK-Y VERIFICATION WORDING ---")
# Print full text of check 3 from each table's verification section
for para_idx in [17, 392, 788]:
    i, style, text = all_text[para_idx]
    # Find check 3
    if 'Verification Check 3' in text:
        start = text.find('Verification Check 3')
        end = text.find('Verification Check 4')
        if end == -1:
            end = start + 400
        snippet = text[start:end].strip()
        print(f"  [P{i}] Check 3 text:")
        print(f"    {snippet!r}")
        print()

print()
print("--- ISSUE 7: TABLE LABELS IN DOCUMENT HEADINGS ---")
print("Report section headings use '2. Table 1', '3. Table 2', '4. Table 3'.")
print("Nowhere in the headings does it say Table A, Table B, or Table E.")
print("However, within the content, the correct table names ARE used:")
py_name_refs = [(i, text) for i, style, text in all_text
                if any(x in text for x in ['Table A', 'Table B', 'Table E', 'TableA', 'TableB', 'TableE'])]
for i, text in py_name_refs[:10]:
    print(f"  [P{i}] {text[:120]!r}")

print()
print("--- ISSUE 8: SECTION 2.4 / 3.4 / 4.4 - IS C++ MENTIONED IN CHECK 4? ---")
for para_idx in [17, 392, 788]:
    i, style, text = all_text[para_idx]
    if 'Verification Check 4' in text:
        start = text.find('Verification Check 4')
        snippet = text[start:start+400].strip()
        print(f"  [P{i}] Check 4 text:")
        print(f"    {snippet!r}")
        print()

print()
print("--- ISSUE 9: DOES CHECK 4 MENTION C++ OR JUST PYTHON? ---")
for para_idx in [17, 392, 788]:
    i, style, text = all_text[para_idx]
    if 'Check 4' in text:
        start = text.find('Check 4')
        chunk = text[start:start+300]
        cpp_mentioned = 'C++' in chunk or 'cpp' in chunk.lower()
        py_mentioned = 'Python' in chunk or '.py' in chunk
        print(f"  [P{i}] C++ mentioned in Check 4: {cpp_mentioned} | Python mentioned: {py_mentioned}")
