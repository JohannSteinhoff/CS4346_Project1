import docx
import sys

doc = docx.Document('/mnt/c/Users/johann/OneDrive/Desktop/Code/Ai_Project_1/Project1-ngq7-Report.docx')
print(f"Number of paragraphs: {len(doc.paragraphs)}")
print(f"Number of tables: {len(doc.tables)}")
print("--- FULL PARAGRAPH TEXT ---")
for i, para in enumerate(doc.paragraphs):
    print(f"[P{i}] STYLE={para.style.name} | {repr(para.text)}")
