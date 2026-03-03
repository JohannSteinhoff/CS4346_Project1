import docx

doc = docx.Document('/mnt/c/Users/johann/OneDrive/Desktop/Code/Ai_Project_1/Project1-ngq7-Report.docx')

all_text = []
for i, para in enumerate(doc.paragraphs):
    all_text.append((i, para.style.name, para.text))

full_text = '\n'.join(t for _, _, t in all_text)

print("=" * 70)
print("AUDIT CHECK 1: PLACEHOLDER / TODO / UNFILLED TEXT")
print("=" * 70)
suspicious = ['[', 'TODO', 'todo', 'PLACEHOLDER', 'placeholder',
              'insert', 'INSERT', 'TBD', 'tbd', 'fill in', 'FILL IN',
              'your name', 'student name']
found_suspicious = []
for i, style, text in all_text:
    for word in suspicious:
        if word in text:
            found_suspicious.append((i, style, word, text))
            break

if found_suspicious:
    for i, style, word, text in found_suspicious:
        print(f"  [P{i}] TRIGGER={word!r} | {text[:200]!r}")
else:
    print("  NONE FOUND - No placeholder text detected.")

print()
print("=" * 70)
print("AUDIT CHECK 2: EMPTY / NEAR-EMPTY PARAGRAPHS AFTER HEADINGS")
print("=" * 70)
for i, style, text in all_text:
    if 'Heading' in style:
        # Check next paragraph
        if i + 1 < len(all_text):
            next_i, next_style, next_text = all_text[i + 1]
            if next_text.strip() == '':
                print(f"  HEADING [P{i}]: {text!r}")
                print(f"    FOLLOWED BY EMPTY PARA [P{next_i}]")

print()
print("=" * 70)
print("AUDIT CHECK 3: TABLES IN DOCX (embedded word tables)")
print("=" * 70)
print(f"  Number of embedded Word tables: {len(doc.tables)}")
for ti, table in enumerate(doc.tables):
    rows = table.rows
    cols = table.columns
    print(f"  Table {ti}: {len(rows)} rows x {len(cols)} cols")
    # Print first row (headers)
    if rows:
        header_cells = [c.text.strip() for c in rows[0].cells]
        print(f"    Headers: {header_cells}")
    # Print last row
    if len(rows) > 1:
        last_cells = [c.text.strip() for c in rows[-1].cells]
        print(f"    Last row: {last_cells}")

print()
print("=" * 70)
print("AUDIT CHECK 4: SECTION HEADINGS PRESENT")
print("=" * 70)
headings = [(i, style, text) for i, style, text in all_text if 'Heading' in style]
for i, style, text in headings:
    print(f"  [P{i}] {style}: {text!r}")

print()
print("=" * 70)
print("AUDIT CHECK 5: PYTHON CODE PRESENCE PER TABLE")
print("=" * 70)
# Find Python code markers
py_markers = [(i, text) for i, style, text in all_text if 'Python Code' in text and '(' in text]
for i, text in py_markers:
    print(f"  [P{i}] {text[:120]!r}")

print()
print("=" * 70)
print("AUDIT CHECK 6: C++ CODE PRESENCE PER TABLE")
print("=" * 70)
cpp_markers = [(i, text) for i, style, text in all_text if 'C++ Code' in text and '(' in text]
for i, text in cpp_markers:
    print(f"  [P{i}] {text[:120]!r}")

print()
print("=" * 70)
print("AUDIT CHECK 7: CODE OUTPUT SECTIONS PER TABLE")
print("=" * 70)
output_markers = [(i, text) for i, style, text in all_text if 'Python Code Output' in text or 'Code Output' in text]
for i, text in output_markers:
    print(f"  [P{i}] {text[:120]!r}")

print()
print("=" * 70)
print("AUDIT CHECK 8: VERIFICATION CHECKS - ALL 4 PER TABLE")
print("=" * 70)
check_markers = [(i, text) for i, style, text in all_text if 'Verification Check' in text]
for i, text in check_markers:
    # Count how many checks are mentioned
    count = text.count('Verification Check')
    print(f"  [P{i}] {count} check(s) mentioned | snippet: {text[:300]!r}")

print()
print("=" * 70)
print("AUDIT CHECK 9: BLANK-Y VERIFICATION (Check 3) PRESENCE")
print("=" * 70)
blank_y = [(i, text) for i, style, text in all_text if 'blank' in text.lower() or 'Blank' in text or 'blank Y' in text.lower() or 'blank-Y' in text.lower() or 'fill in Y' in text.lower() or 'filled in Y' in text.lower() or 'filled in' in text.lower()]
for i, text in blank_y:
    print(f"  [P{i}] {text[:200]!r}")

print()
print("=" * 70)
print("AUDIT CHECK 10: PROJECT CONCLUSIONS SECTION")
print("=" * 70)
proj_conc = [(i, style, text) for i, style, text in all_text if 'Project Conclusion' in text or ('Conclusion' in text and 'Heading 1' in style)]
for i, style, text in proj_conc:
    print(f"  [P{i}] STYLE={style} | {text[:300]!r}")
# Also show what follows
for i, style, text in proj_conc:
    if 'Heading' in style and i+1 < len(all_text):
        ni, ns, nt = all_text[i+1]
        print(f"  -> NEXT PARA [P{ni}]: {nt[:200]!r}")

print()
print("=" * 70)
print("AUDIT CHECK 11: PSR DEFINITION SECTION")
print("=" * 70)
psr_def = [(i, style, text) for i, style, text in all_text if 'PSR Definition' in text or ('PSR' in text and 'Definition' in text)]
for i, style, text in psr_def:
    print(f"  [P{i}] STYLE={style} | {text[:300]!r}")

print()
print("=" * 70)
print("AUDIT CHECK 12: DUPLICATE PYTHON CODE FOR TABLE E")
print("=" * 70)
tableE_py = [(i, text) for i, style, text in all_text if 'TableE.py' in text or ('Table E' in text and 'Python' in text)]
for i, text in tableE_py:
    print(f"  [P{i}] {text[:150]!r}")

print()
print("=" * 70)
print("AUDIT CHECK 13: C++ OUTPUT SECTIONS (if any)")
print("=" * 70)
cpp_output = [(i, text) for i, style, text in all_text if 'C++ Code Output' in text or 'C++ Output' in text or 'cpp output' in text.lower()]
if cpp_output:
    for i, text in cpp_output:
        print(f"  [P{i}] {text[:150]!r}")
else:
    print("  WARNING: No C++ code output sections found in the report.")

print()
print("=" * 70)
print("AUDIT CHECK 14: TABLE A, B, E NAMING (confirm which table is which)")
print("=" * 70)
table_refs = [(i, style, text) for i, style, text in all_text
              if any(x in text for x in ['Table A', 'Table B', 'Table E', 'TableA', 'TableB', 'TableE'])
              and 'Heading' in style]
for i, style, text in table_refs:
    print(f"  [P{i}] {style}: {text!r}")
# Also check section 1/2/3 vs A/B/E naming
heading1s = [(i, text) for i, style, text in all_text if style == 'Heading 1']
print(f"  All Heading 1s: {[(i, t) for i, t in heading1s]}")

print()
print("=" * 70)
print("AUDIT CHECK 15: CHECK 4 - PYTHON/C++ CODE RAN AND MATCHES (copy of code)")
print("=" * 70)
check4 = [(i, text) for i, style, text in all_text if 'Check 4' in text or 'Verification Check 4' in text]
for i, text in check4:
    print(f"  [P{i}] {text[:400]!r}")

print()
print("=" * 70)
print("AUDIT CHECK 16: ARE ACTUAL DATA TABLES (Word tables) FILLED OR EMPTY?")
print("=" * 70)
for ti, table in enumerate(doc.tables):
    rows = table.rows
    empty_cells = 0
    total_cells = 0
    for row in rows:
        for cell in row.cells:
            total_cells += 1
            if cell.text.strip() == '':
                empty_cells += 1
    print(f"  Table {ti}: {len(rows)} rows, {total_cells} total cells, {empty_cells} empty cells")
    # Sample a few rows
    sample_rows = min(3, len(rows))
    for ri in range(sample_rows):
        cells = [c.text.strip()[:15] for c in rows[ri].cells]
        print(f"    Row {ri}: {cells}")

print()
print("=" * 70)
print("AUDIT CHECK 17: SUMMARY OF WHAT IS PRESENT/MISSING")
print("=" * 70)

# Check structure
sections_found = {}
for i, style, text in all_text:
    if style == 'Heading 1':
        sections_found[text] = i

print("  Top-level sections found:")
for s, i in sections_found.items():
    print(f"    [P{i}] {s!r}")

# Check for all 4 verification checks in each table section
import re
table_sections = {}
current_section = None
current_paras = []
for i, style, text in all_text:
    if style == 'Heading 1':
        if current_section:
            table_sections[current_section] = current_paras
        current_section = text
        current_paras = []
    else:
        current_paras.append((i, text))
if current_section:
    table_sections[current_section] = current_paras

print()
print("  Per-section verification check analysis:")
for section, paras in table_sections.items():
    combined = ' '.join(t for _, t in paras)
    checks = {}
    for c in range(1, 5):
        checks[c] = f'Check {c}' in combined or f'Verification Check {c}' in combined
    py_present = 'def predict' in combined or 'import' in combined or 'data = [' in combined or '# CS 4346' in combined
    cpp_present = '#include' in combined or 'using namespace std' in combined
    py_output = 'Python Code Output' in combined or ('Accuracy' in combined and '100%' in combined)
    has_rules = 'Rule 1:' in combined or 'IF ' in combined
    print(f"  Section: {section!r}")
    print(f"    Check 1 (rules satisfy all rows): {checks[1]}")
    print(f"    Check 2 (too many rules?):        {checks[2]}")
    print(f"    Check 3 (blank-Y verification):   {checks[3]}")
    print(f"    Check 4 (code matches ChatGPT):   {checks[4]}")
    print(f"    Python code present:               {py_present}")
    print(f"    C++ code present:                  {cpp_present}")
    print(f"    Python output present:             {py_output}")
    print(f"    Rules present:                     {has_rules}")
    print()
